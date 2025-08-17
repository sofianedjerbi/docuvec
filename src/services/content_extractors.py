"""Content extraction methods for different MIME types"""

import io
import json
import base64
from typing import Optional, Dict, Any, Tuple
from pathlib import Path

from pypdf import PdfReader
import trafilatura

from src.core.logger import setup_logger
from src.services.html_extractor import HTMLExtractor


class ContentExtractors:
    """Content extraction methods for various file types"""
    
    def __init__(self, enable_ocr: bool = False):
        """
        Initialize content extractors
        
        Args:
            enable_ocr: Whether to enable OCR for images and scanned PDFs
        """
        self.enable_ocr = enable_ocr
        self.logger = setup_logger(self.__class__.__name__)
        self.html_extractor = HTMLExtractor()
        
        # Lazy load optional libraries
        self._docx = None
        self._mammoth = None
        self._pptx = None
        self._pandas = None
        self._openpyxl = None
        self._xlrd = None
        self._pytesseract = None
        self._PIL = None
        self._pdfplumber = None
        
        # Store metadata for last extraction
        self.last_metadata = {}
    
    def extract_html(self, content: bytes, url: str = "") -> Optional[str]:
        """Extract text from HTML content using advanced tiered extraction"""
        try:
            # Decode bytes to string
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            
            # Use the advanced HTML extractor with tiered fallback
            text, metadata = self.html_extractor.extract(content, url)
            
            # Store metadata for later use
            self.last_metadata = metadata
            
            # Post-process the text
            if text:
                text = self.html_extractor.post_process(text)
            
            # Log extraction method used
            if metadata.get('extraction_method'):
                self.logger.debug(f"Extracted HTML using {metadata['extraction_method']}")
            
            return text
            
        except Exception as e:
            self.logger.error(f"Failed to extract HTML content: {e}")
            return None
    
    def extract_pdf(self, content: bytes, enable_ocr_fallback: bool = False) -> Optional[str]:
        """Extract text from PDF content"""
        try:
            # Try pypdf first (fast, works for most PDFs)
            reader = PdfReader(io.BytesIO(content))
            pages = []
            
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text)
            
            full_text = "\n\n".join(pages)
            
            # If no text extracted and OCR is enabled, try OCR
            if not full_text.strip() and enable_ocr_fallback and self.enable_ocr:
                self.logger.info("No text in PDF, attempting OCR extraction")
                full_text = self._extract_pdf_with_ocr(content)
            
            # If still no text, try pdfplumber as fallback
            if not full_text.strip():
                full_text = self._extract_pdf_with_pdfplumber(content)
            
            return full_text if full_text.strip() else None
            
        except Exception as e:
            self.logger.error(f"Failed to extract PDF content: {e}")
            return None
    
    def _extract_pdf_with_pdfplumber(self, content: bytes) -> Optional[str]:
        """Extract PDF text using pdfplumber (better for complex layouts)"""
        try:
            if self._pdfplumber is None:
                import pdfplumber
                self._pdfplumber = pdfplumber
            
            pages = []
            with self._pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            
            return "\n\n".join(pages)
            
        except ImportError:
            self.logger.warning("pdfplumber not installed, skipping advanced PDF extraction")
            return None
        except Exception as e:
            self.logger.error(f"Failed to extract PDF with pdfplumber: {e}")
            return None
    
    def _extract_pdf_with_ocr(self, content: bytes) -> Optional[str]:
        """Extract text from scanned PDF using OCR"""
        try:
            if self._pytesseract is None:
                import pytesseract
                self._pytesseract = pytesseract
            
            if self._PIL is None:
                from PIL import Image
                self._PIL = Image
            
            # Convert PDF to images and OCR each page
            # This requires pdf2image library
            try:
                from pdf2image import convert_from_bytes
                
                images = convert_from_bytes(content)
                texts = []
                
                for i, image in enumerate(images):
                    self.logger.info(f"Running OCR on page {i+1}/{len(images)}")
                    text = self._pytesseract.image_to_string(image)
                    if text.strip():
                        texts.append(text)
                
                return "\n\n".join(texts)
                
            except ImportError:
                self.logger.warning("pdf2image not installed, cannot OCR PDF files")
                return None
                
        except ImportError:
            self.logger.warning("pytesseract/PIL not installed, OCR disabled")
            return None
        except Exception as e:
            self.logger.error(f"Failed to OCR PDF: {e}")
            return None
    
    def extract_docx(self, content: bytes) -> Optional[str]:
        """Extract text from DOCX files"""
        try:
            # Try python-docx first
            if self._docx is None:
                import docx
                self._docx = docx
            
            doc = self._docx.Document(io.BytesIO(content))
            paragraphs = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
            
        except ImportError:
            # Fallback to mammoth
            return self._extract_docx_with_mammoth(content)
        except Exception as e:
            self.logger.error(f"Failed to extract DOCX: {e}")
            return None
    
    def _extract_docx_with_mammoth(self, content: bytes) -> Optional[str]:
        """Extract DOCX using mammoth (preserves more formatting)"""
        try:
            if self._mammoth is None:
                import mammoth
                self._mammoth = mammoth
            
            result = self._mammoth.extract_raw_text(io.BytesIO(content))
            return result.value
            
        except ImportError:
            self.logger.warning("mammoth not installed, cannot extract DOCX")
            return None
        except Exception as e:
            self.logger.error(f"Failed to extract DOCX with mammoth: {e}")
            return None
    
    def extract_doc(self, content: bytes) -> Optional[str]:
        """Extract text from legacy DOC files"""
        # For legacy DOC files, we would need python-docx2txt or similar
        # For now, try to extract as DOCX (some work)
        return self.extract_docx(content)
    
    def extract_pptx(self, content: bytes) -> Optional[str]:
        """Extract text from PowerPoint files"""
        try:
            if self._pptx is None:
                from pptx import Presentation
                self._pptx = Presentation
            
            prs = self._pptx(io.BytesIO(content))
            texts = []
            
            for slide in prs.slides:
                slide_texts = []
                for shape in slide.shapes:
                    if hasattr(shape, "text") and shape.text.strip():
                        slide_texts.append(shape.text)
                
                if slide_texts:
                    texts.append("\n".join(slide_texts))
            
            return "\n\n---\n\n".join(texts)  # Separate slides
            
        except ImportError:
            self.logger.warning("python-pptx not installed, cannot extract PPTX")
            return None
        except Exception as e:
            self.logger.error(f"Failed to extract PPTX: {e}")
            return None
    
    def extract_ppt(self, content: bytes) -> Optional[str]:
        """Extract text from legacy PPT files"""
        # For legacy PPT, we'd need additional libraries
        # Try as PPTX (some work)
        return self.extract_pptx(content)
    
    def extract_xlsx(self, content: bytes) -> Optional[str]:
        """Extract text from Excel files"""
        try:
            if self._pandas is None:
                import pandas as pd
                self._pandas = pd
            
            # Read all sheets
            excel_file = io.BytesIO(content)
            sheets = self._pandas.read_excel(excel_file, sheet_name=None)
            
            texts = []
            for sheet_name, df in sheets.items():
                texts.append(f"## Sheet: {sheet_name}\n")
                # Convert dataframe to string
                texts.append(df.to_string())
            
            return "\n\n".join(texts)
            
        except ImportError:
            self.logger.warning("pandas not installed, cannot extract XLSX")
            return None
        except Exception as e:
            self.logger.error(f"Failed to extract XLSX: {e}")
            return None
    
    def extract_xls(self, content: bytes) -> Optional[str]:
        """Extract text from legacy Excel files"""
        try:
            if self._pandas is None:
                import pandas as pd
                self._pandas = pd
            
            # For legacy XLS, pandas with xlrd backend
            excel_file = io.BytesIO(content)
            sheets = self._pandas.read_excel(excel_file, sheet_name=None, engine='xlrd')
            
            texts = []
            for sheet_name, df in sheets.items():
                texts.append(f"## Sheet: {sheet_name}\n")
                texts.append(df.to_string())
            
            return "\n\n".join(texts)
            
        except ImportError:
            self.logger.warning("pandas/xlrd not installed, cannot extract XLS")
            return None
        except Exception as e:
            self.logger.error(f"Failed to extract XLS: {e}")
            return None
    
    def extract_text(self, content: bytes) -> Optional[str]:
        """Extract plain text"""
        try:
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            self.logger.error(f"Failed to decode text: {e}")
            return None
    
    def extract_markdown(self, content: bytes) -> Optional[str]:
        """Extract markdown (just decode, already in good format)"""
        return self.extract_text(content)
    
    def extract_json(self, content: bytes) -> Optional[str]:
        """Extract and format JSON content"""
        try:
            data = json.loads(content.decode('utf-8'))
            # Convert to readable text format
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            self.logger.error(f"Failed to parse JSON: {e}")
            return self.extract_text(content)  # Fallback to plain text
    
    def extract_csv(self, content: bytes) -> Optional[str]:
        """Extract CSV content"""
        try:
            if self._pandas is None:
                import pandas as pd
                self._pandas = pd
            
            df = self._pandas.read_csv(io.BytesIO(content))
            return df.to_string()
            
        except ImportError:
            # Fallback to plain text
            return self.extract_text(content)
        except Exception as e:
            self.logger.error(f"Failed to parse CSV: {e}")
            return self.extract_text(content)
    
    def extract_image_ocr(self, content: bytes) -> Optional[str]:
        """Extract text from images using OCR"""
        if not self.enable_ocr:
            self.logger.info("OCR disabled, skipping image")
            return None
        
        try:
            if self._pytesseract is None:
                import pytesseract
                self._pytesseract = pytesseract
            
            if self._PIL is None:
                from PIL import Image
                self._PIL = Image
            
            image = self._PIL.open(io.BytesIO(content))
            text = self._pytesseract.image_to_string(image)
            
            return text if text.strip() else None
            
        except ImportError:
            self.logger.warning("pytesseract/PIL not installed, cannot OCR images")
            return None
        except Exception as e:
            self.logger.error(f"Failed to OCR image: {e}")
            return None
    
    def skip_image(self, content: bytes) -> Optional[str]:
        """Skip image processing when OCR is disabled"""
        self.logger.debug("Skipping image (OCR disabled)")
        return None
    
    def extract_fallback(self, content: bytes) -> Optional[str]:
        """Fallback extraction - try as text"""
        self.logger.warning("Using fallback text extraction for unknown content type")
        return self.extract_text(content)